"""
Generates small, real (non-mocked) documents for Module 6 parser tests —
using actual libraries (fpdf2, python-docx, openpyxl, Pillow) rather than
hand-written bytes, so adapters are exercised against genuine file formats.
"""

import io

import docx
import openpyxl
from fpdf import FPDF
from PIL import Image


def make_pdf_bytes(text: str = "Invoice #1042") -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=14)
    pdf.cell(0, 10, text, new_x="LMARGIN", new_y="NEXT")
    with pdf.table() as table:
        row = table.row()
        row.cell("Item")
        row.cell("Qty")
        row2 = table.row()
        row2.cell("Widget")
        row2.cell("5")
    return bytes(pdf.output())


def make_corrupt_pdf_bytes() -> bytes:
    return b"%PDF-1.4\nthis is not a valid pdf body"


def make_docx_bytes(text: str = "Invoice #1042", with_image: bool = True) -> bytes:
    document = docx.Document()
    document.add_paragraph(text)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Qty"
    table.cell(1, 0).text = "Widget"
    table.cell(1, 1).text = "5"

    if with_image:
        image_buffer = io.BytesIO()
        Image.new("RGB", (30, 20), color="blue").save(image_buffer, format="PNG")
        image_buffer.seek(0)
        document.add_picture(image_buffer)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_xlsx_bytes() -> bytes:
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet["A1"] = "Name"
    worksheet["B1"] = "Qty"
    worksheet["A2"] = "Widget"
    worksheet["B2"] = 5
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def make_png_bytes(width: int = 100, height: int = 50) -> bytes:
    buffer = io.BytesIO()
    Image.new("RGB", (width, height), color="red").save(buffer, format="PNG")
    return buffer.getvalue()
